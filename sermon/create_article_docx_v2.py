# -*- coding: utf-8 -*-
"""
4차원영성 생각 아티클 v2 — 보강/보충판
- 초신자 대상 쉬운 설명
- 가독성 높은 디자인 (색상, 글박스)
- 페이지 넘버링
- 각주 처리 (설교 참조)
- 풍성한 내용 (예화, 간증)
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

doc = Document()

# ═══════════════════════════════════════════════════════════
# 색상 팔레트
# ═══════════════════════════════════════════════════════════
C_NAVY    = RGBColor(0x1B, 0x4F, 0x72)
C_TEAL    = RGBColor(0x11, 0x7A, 0x65)
C_GOLD    = RGBColor(0xB7, 0x95, 0x0B)
C_RED     = RGBColor(0x8B, 0x00, 0x00)
C_BODY    = RGBColor(0x33, 0x33, 0x33)
C_SUBTLE  = RGBColor(0x66, 0x66, 0x66)
C_LIGHT   = RGBColor(0x99, 0x99, 0x99)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_ACCENT  = RGBColor(0x2E, 0x86, 0xC1)

# ═══════════════════════════════════════════════════════════
# 페이지 설정
# ═══════════════════════════════════════════════════════════
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ═══════════════════════════════════════════════════════════
# 스타일 설정
# ═══════════════════════════════════════════════════════════
style_normal = doc.styles['Normal']
style_normal.font.name = '맑은 고딕'
style_normal.font.size = Pt(10.5)
style_normal.font.color.rgb = C_BODY
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.6

for level in [1, 2, 3]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '맑은 고딕'
    if level == 1:
        hs.font.size = Pt(20)
        hs.font.bold = True
        hs.font.color.rgb = C_NAVY
        hs.paragraph_format.space_before = Pt(36)
        hs.paragraph_format.space_after = Pt(6)
    elif level == 2:
        hs.font.size = Pt(14)
        hs.font.bold = True
        hs.font.color.rgb = C_TEAL
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(8)
    elif level == 3:
        hs.font.size = Pt(12)
        hs.font.bold = True
        hs.font.color.rgb = C_NAVY
        hs.paragraph_format.space_before = Pt(10)
        hs.paragraph_format.space_after = Pt(6)

# ═══════════════════════════════════════════════════════════
# 각주(Footnote) 시스템
# ═══════════════════════════════════════════════════════════
class FootnoteManager:
    def __init__(self, document):
        self.doc = document
        self.next_id = 1
        self._init_part()

    def _init_part(self):
        nsmap_fn = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        }
        self.root = etree.Element(qn('w:footnotes'), nsmap=nsmap_fn)

        # separator (id=-1)
        sep = etree.SubElement(self.root, qn('w:footnote'))
        sep.set(qn('w:type'), 'separator')
        sep.set(qn('w:id'), '-1')
        sp = etree.SubElement(sep, qn('w:p'))
        sr = etree.SubElement(sp, qn('w:r'))
        etree.SubElement(sr, qn('w:separator'))

        # continuation (id=0)
        cont = etree.SubElement(self.root, qn('w:footnote'))
        cont.set(qn('w:type'), 'continuationSeparator')
        cont.set(qn('w:id'), '0')
        cp = etree.SubElement(cont, qn('w:p'))
        cr = etree.SubElement(cp, qn('w:r'))
        etree.SubElement(cr, qn('w:continuationSeparator'))

        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        blob = etree.tostring(self.root, xml_declaration=True, encoding='UTF-8', standalone=True)
        self.part = Part(
            PackURI('/word/footnotes.xml'),
            'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml',
            blob,
            self.doc.part.package
        )
        self.doc.part.relate_to(
            self.part,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
        )

    def add(self, paragraph, text):
        fid = self.next_id
        self.next_id += 1

        # --- footnotes.xml에 각주 추가 ---
        fn = etree.SubElement(self.root, qn('w:footnote'))
        fn.set(qn('w:id'), str(fid))

        fn_p = etree.SubElement(fn, qn('w:p'))
        fn_pPr = etree.SubElement(fn_p, qn('w:pPr'))
        fn_pStyle = etree.SubElement(fn_pPr, qn('w:pStyle'))
        fn_pStyle.set(qn('w:val'), 'FootnoteText')

        # 각주 참조 마크
        fn_r1 = etree.SubElement(fn_p, qn('w:r'))
        fn_rPr1 = etree.SubElement(fn_r1, qn('w:rPr'))
        fn_rStyle1 = etree.SubElement(fn_rPr1, qn('w:rStyle'))
        fn_rStyle1.set(qn('w:val'), 'FootnoteReference')
        etree.SubElement(fn_r1, qn('w:footnoteRef'))

        # 공백
        fn_rs = etree.SubElement(fn_p, qn('w:r'))
        fn_ts = etree.SubElement(fn_rs, qn('w:t'))
        fn_ts.set(qn('xml:space'), 'preserve')
        fn_ts.text = ' '

        # 각주 본문
        fn_r2 = etree.SubElement(fn_p, qn('w:r'))
        fn_rPr2 = etree.SubElement(fn_r2, qn('w:rPr'))
        fn_sz = etree.SubElement(fn_rPr2, qn('w:sz'))
        fn_sz.set(qn('w:val'), '18')
        fn_t = etree.SubElement(fn_r2, qn('w:t'))
        fn_t.set(qn('xml:space'), 'preserve')
        fn_t.text = text

        self.part._blob = etree.tostring(self.root, xml_declaration=True, encoding='UTF-8', standalone=True)

        # --- 본문에 각주 참조 삽입 ---
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'FootnoteReference')
        rPr.append(rStyle)
        run.append(rPr)
        fnRef = OxmlElement('w:footnoteReference')
        fnRef.set(qn('w:id'), str(fid))
        run.append(fnRef)
        paragraph._p.append(run)
        return fid

fn_mgr = FootnoteManager(doc)

# ═══════════════════════════════════════════════════════════
# 페이지 번호 (푸터)
# ═══════════════════════════════════════════════════════════
def add_page_numbers():
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = p.add_run()
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fc1)

    run2 = p.add_run()
    it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve')
    it.text = ' PAGE '
    run2._r.append(it)

    run3 = p.add_run()
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fc2)

    for r in p.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = C_LIGHT

add_page_numbers()

# ═══════════════════════════════════════════════════════════
# 헬퍼 함수들
# ═══════════════════════════════════════════════════════════
def P(text, bold=False, italic=False, size=None, color=None, align=None,
      space_before=0, space_after=6, indent=0):
    """일반 단락 추가"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    return p

def P_multi(parts, space_before=0, space_after=6, indent=0):
    """여러 스타일이 섞인 단락. parts = [(text, {bold, italic, color, size}), ...]"""
    p = doc.add_paragraph()
    for text, fmt in parts:
        run = p.add_run(text)
        run.bold = fmt.get('bold', False)
        run.italic = fmt.get('italic', False)
        if 'color' in fmt: run.font.color.rgb = fmt['color']
        if 'size' in fmt: run.font.size = Pt(fmt['size'])
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    return p

def quote_box(text, bg='EBF5FB', border='2E86C1'):
    """성경 말씀 인용 박스 (좌측 파란 라인 + 배경)"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    # 셀 배경색
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), bg)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)
    # 좌측 두꺼운 테두리
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:start')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:color'), border)
    left.set(qn('w:space'), '0')
    borders.append(left)
    for side in ['top', 'bottom', 'end']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:color'), 'auto')
        el.set(qn('w:space'), '0')
        borders.append(el)
    tcPr.append(borders)
    # 셀 여백
    mar = OxmlElement('w:tcMar')
    for side_name, side_tag in [('top','top'),('bottom','bottom'),('left','start'),('right','end')]:
        m = OxmlElement(f'w:{side_tag}')
        m.set(qn('w:w'), '120')
        m.set(qn('w:type'), 'dxa')
        mar.append(m)
    tcPr.append(mar)

    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1a, 0x5a, 0x7a)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    return tbl

def key_box(title, text, bg='FEF9E7', border='F39C12'):
    """핵심 포인트 박스 (금색 테두리)"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), bg)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)
    borders = OxmlElement('w:tcBorders')
    for side in ['start', 'top', 'end', 'bottom']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '8')
        el.set(qn('w:color'), border)
        el.set(qn('w:space'), '0')
        borders.append(el)
    tcPr.append(borders)
    mar = OxmlElement('w:tcMar')
    for _, tag in [('t','top'),('b','bottom'),('l','start'),('r','end')]:
        m = OxmlElement(f'w:{tag}')
        m.set(qn('w:w'), '140')
        m.set(qn('w:type'), 'dxa')
        mar.append(m)
    tcPr.append(mar)

    p = cell.paragraphs[0]
    run_t = p.add_run(f'{title}\n')
    run_t.bold = True
    run_t.font.size = Pt(10)
    run_t.font.color.rgb = RGBColor(0xB7, 0x95, 0x0B)
    run_b = p.add_run(text)
    run_b.font.size = Pt(10)
    run_b.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    return tbl

def tip_box(text, bg='EAFAF1', border='27AE60'):
    """실천 TIP 박스 (초록색)"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), bg)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)
    borders = OxmlElement('w:tcBorders')
    for side in ['start', 'top', 'end', 'bottom']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '8')
        el.set(qn('w:color'), border)
        el.set(qn('w:space'), '0')
        borders.append(el)
    tcPr.append(borders)
    mar = OxmlElement('w:tcMar')
    for _, tag in [('t','top'),('b','bottom'),('l','start'),('r','end')]:
        m = OxmlElement(f'w:{tag}')
        m.set(qn('w:w'), '140')
        m.set(qn('w:type'), 'dxa')
        mar.append(m)
    tcPr.append(mar)

    p = cell.paragraphs[0]
    run_t = p.add_run('실천 TIP\n')
    run_t.bold = True
    run_t.font.size = Pt(10)
    run_t.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    run_b = p.add_run(text)
    run_b.font.size = Pt(10)
    run_b.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    return tbl

def term_box(term, definition, bg='F2F3F4', border='7F8C8D'):
    """용어 설명 박스 (회색)"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), bg)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)
    borders = OxmlElement('w:tcBorders')
    for side in ['start', 'top', 'end', 'bottom']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:color'), border)
        el.set(qn('w:space'), '0')
        borders.append(el)
    tcPr.append(borders)
    mar = OxmlElement('w:tcMar')
    for _, tag in [('t','top'),('b','bottom'),('l','start'),('r','end')]:
        m = OxmlElement(f'w:{tag}')
        m.set(qn('w:w'), '140')
        m.set(qn('w:type'), 'dxa')
        mar.append(m)
    tcPr.append(mar)

    p = cell.paragraphs[0]
    rt = p.add_run(f'{term}\n')
    rt.bold = True
    rt.font.size = Pt(10)
    rt.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    rd = p.add_run(definition)
    rd.font.size = Pt(9.5)
    rd.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    return tbl

def sub(text):
    """▌ 소제목"""
    p = doc.add_paragraph()
    run = p.add_run(f'▌ {text}')
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = C_NAVY
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text):
    """• 불릿"""
    p = doc.add_paragraph()
    run = p.add_run(f'•  {text}')
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    return p

def numbered(num, text):
    """① ② ③ ..."""
    c = '①②③④⑤⑥⑦⑧⑨⑩'
    p = doc.add_paragraph()
    run = p.add_run(f'{c[num-1]} {text}')
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = C_NAVY
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p

def foundation(text):
    """✦ 기초 섹션"""
    p = doc.add_paragraph()
    run = p.add_run(f'✦  {text}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = C_RED
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p

def blank(pts=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pts)
    p.paragraph_format.space_after = Pt(0)

def divider():
    """구분선"""
    p = doc.add_paragraph()
    run = p.add_run('─' * 50)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


# ═══════════════════════════════════════════════════════════════
#                          본 문 시 작
# ═══════════════════════════════════════════════════════════════

# ── 표지 ──
blank(20)
P('4차원의 영성', bold=True, size=28, color=C_NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
P("'생각'의 능력", bold=True, size=16, color=C_TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
divider()
P('조용기 목사 설교(1981~2020년) 391편 분석 기반', size=10, color=C_SUBTLE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
P('대상: 일반 성도 · 초신자', size=10, color=C_SUBTLE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 들어가며
# ══════════════════════════════════════════════════════════
doc.add_heading('들어가며', level=1)
doc.add_heading('생각이 인생을 만든다', level=2)

P('여러분은 오늘 아침에 어떤 생각을 하며 하루를 시작하셨습니까?', bold=True, size=11, color=C_NAVY, space_before=8)
blank()

P('"오늘도 힘든 하루가 되겠지..." 라고 생각하셨습니까? '
  '아니면 "오늘 하나님이 좋은 일을 예비해 두셨을 거야" 라고 생각하셨습니까? '
  '이 작은 차이가, 실은 삶 전체를 바꾸는 거대한 차이입니다.')

quote_box('"지킬만한 것보다 더욱 네 마음을 지키라. 생명의 근원이 이에서 남이니라." (잠언 4:23)')
blank()

P('성경이 재산도, 건강도, 인간관계도 아닌 \'마음\'을 가장 먼저 지키라고 말씀하는 이유는 무엇일까요? '
  '그것은 우리의 생각이 단순히 머릿속에서 지나가는 관념이 아니라, '
  '삶의 방향과 결과를 실질적으로 결정하는 \'생명의 근원\'이기 때문입니다.')

P('이 진리를 조용기 목사는 약 40년간의 사역에서 수천 번 반복해서 가르쳤습니다. '
  '그것이 바로 \'4차원의 영성\'입니다.')
blank()

# 4차원 영성 용어 설명 박스
term_box(
    '4차원의 영성이란?',
    '우리가 사는 눈에 보이는 세계(3차원)는 눈에 보이지 않는 영적 세계(4차원)의 지배를 받습니다. '
    '마치 컴퓨터 화면에 보이는 것(3차원)이 보이지 않는 프로그램 코드(4차원)에 의해 결정되는 것과 같습니다. '
    '그리고 이 영적 차원을 움직이는 네 가지 도구가 있습니다:\n'
    '① 생각 — 마음속에서 무엇을 떠올리고 품느냐\n'
    '② 꿈 — 어떤 미래를 그리느냐\n'
    '③ 믿음 — 보이지 않는 것을 실체로 받아들이느냐\n'
    '④ 말 — 입으로 무엇을 선포하느냐\n'
    '이 네 가지 중 첫 번째이자 가장 근본적인 것이 바로 \'생각\'입니다.'
)
blank()

P('생각이 바뀌지 않으면 꿈도, 믿음도, 말도 바뀌지 않습니다. '
  '생각은 모든 변화의 출발점입니다. '
  '이 문서는 조용기 목사의 설교 391편에서 추출한 원리를 바탕으로, '
  '생각이 왜 중요한지, 왜 부정적이 되는지, 그리고 어떻게 바꿀 수 있는지를 '
  '누구나 이해할 수 있도록 정리합니다.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1부
# ══════════════════════════════════════════════════════════
doc.add_heading('1부', level=1)
doc.add_heading('생각은 왜 중요한가', level=2)

# ① 생각이 곧 그 사람의 실체이다
numbered(1, '생각이 곧 그 사람의 실체이다')
quote_box('"그 마음의 생각이 어떠하면 그 사람도 그러하니." (잠언 23:7)')
blank()

P('성경은 이 원리를 명확하게 선언합니다. 겉으로 보이는 외모, 직업, 사회적 지위가 '
  '그 사람의 진정한 모습이 아닙니다. 그 사람의 마음속에 어떤 생각이 자리잡고 있느냐가 '
  '그 사람의 진정한 모습입니다.')

p = P('인간이 지구의 주인이 된 것은 체력 때문이 아닙니다. 사자보다 힘이 약하고, '
  '독수리처럼 날 수도 없으며, 물고기처럼 물속에서 살 수도 없습니다. '
  '그런데도 인간이 이 지구를 정복한 것은 오직 \'생각할 수 있는 능력\' 때문입니다. '
  '오늘날 과학 만능의 시대를 만든 것도 생각의 산물이며, '
  '동시에 인류 파멸의 위기까지 몰아넣은 것도 인간 생각의 결과입니다.')
fn_mgr.add(p, '조용기 목사, "생각을 바꿔라" (1984.11.11) 설교 참조')
blank()

P('같은 가게에서 일하는 두 직원을 떠올려 보십시오. 한 사람은 "이 일은 보잘것없는 일이야, '
  '나는 여기서 벗어나야 해"라고 생각합니다. 다른 사람은 "이 일을 통해 배우고 성장하자, '
  '하나님이 여기에 나를 두신 이유가 있을 거야"라고 생각합니다. '
  '1년 뒤, 두 사람의 삶은 완전히 다른 방향으로 흘러갑니다. '
  '환경은 같았습니다. 달라진 것은 \'생각\'이었습니다.')

# ② 생각은 씨앗처럼 반드시 열매를 맺는다
numbered(2, '생각은 씨앗처럼 반드시 열매를 맺는다')
P('예수님은 씨 뿌리는 자의 비유(마태복음 13장)를 통해 마음과 생각의 관계를 설명하셨습니다. '
  '네 종류의 밭 — 길가, 자갈밭, 가시밭, 옥토 — 은 바로 사람의 생각의 상태를 보여줍니다.')

term_box(
    '씨 뿌리는 자의 비유 (마태복음 13:3-8)',
    '① 길가 — 굳어진 마음. 말씀이 들어와도 금방 빼앗김\n'
    '② 자갈밭 — 깊이 없는 마음. 처음엔 기뻐하나 시련이 오면 포기\n'
    '③ 가시밭 — 세상 걱정과 재물의 유혹으로 말씀이 자라지 못함\n'
    '④ 좋은 밭 — 말씀을 듣고 깨달아 열매를 맺는 마음'
)
blank()

P('매일의 생각이 곧 매일 마음의 밭에 뿌리는 씨앗입니다. '
  '"나는 못한다, 나는 안 된다, 나는 실패자다" — 이런 부정적 생각의 씨앗을 심으면 '
  '반드시 파멸의 열매를 거두게 됩니다. 반대로 "하나님이 나와 함께하신다, 나는 할 수 있다, '
  '나는 축복받았다" — 이런 생각의 씨앗은 성공과 축복의 열매를 맺습니다.')
blank()

p = P('조용기 목사는 이 원리를 농부의 비유로 설명합니다. '
  '"행복은 농부가 농사를 짓는 것과 똑같습니다. 농부가 밭을 가꾸고 씨를 뿌려서 잡초를 제거하고 '
  '가꾸어 농사를 지어 열매 맺는 것처럼, 행복이란 것은 반드시 여러분 생활 속에 '
  '행복의 씨앗을 심어가지고 행복의 열매를 거두는 것입니다."')
fn_mgr.add(p, '조용기 목사, "행복과 마음의 작정" (1981.1.11)')

# ③ 생각이 행동을 결정하고, 행동이 인생을 결정한다
numbered(3, '생각이 행동을 결정하고, 행동이 인생을 결정한다')
P('생각은 눈에 보이지 않지만, 결국 눈에 보이는 결과로 나타납니다. '
  '그 과정은 이렇습니다:')

key_box('생각의 연쇄 반응',
        '생각 → 감정 → 태도 → 행동 → 습관 → 인격 → 운명\n\n'
        '\'나는 안된다\'고 생각하면 → 불안한 감정 → 소극적 태도 → 시도하지 않음 → 포기의 습관\n'
        '\'나는 할 수 있다\'고 생각하면 → 용기 → 적극적 태도 → 도전 → 성공의 열매')
blank()

P('같은 환경, 같은 출발점에서도 생각의 방향에 따라 한 사람은 성공하고 '
  '다른 사람은 실패합니다. 이것은 운명이 아닙니다. 생각의 선택입니다.')
blank()

p = P('성경에도 이 원리를 보여주는 대표적인 이야기가 있습니다. '
  '모세가 이스라엘 백성을 이끌고 가나안 땅 앞에 이르렀을 때, 열두 명의 정탐꾼을 보냈습니다. '
  '같은 땅을 같은 기간 동안 보고 왔는데, 열 명은 이렇게 보고했습니다. '
  '"그 땅의 거인족 앞에서 우리는 메뚜기 같았습니다." '
  '그러나 여호수아와 갈렙 두 사람은 달랐습니다. '
  '"하나님이 우리와 함께하시니, 저들은 우리의 밥입니다!" '
  '같은 상황을 보고도 생각이 달랐고, 그 결과 오직 여호수아와 갈렙만이 약속의 땅에 들어갔습니다.')
fn_mgr.add(p, '민수기 13-14장; "생각을 바꿔라" (1984.11.11) 설교에서 상세히 다룸')

# ④ 마음에는 두 가지 눈이 있다
numbered(4, '마음에는 두 가지 눈이 있다')
quote_box('"너희 마음의 눈을 밝히사 그의 부르심의 소망이 무엇이며." (에베소서 1:18)')
blank()

P('사람에게는 두 가지 눈이 있습니다. 육신의 눈과 마음의 눈입니다.')
blank()

P('육신의 눈은 보이는 세계를 봅니다. 현실의 어려움, 통장 잔고, 병원 진단서, '
  '험난한 세상 — 이것이 육신의 눈에 보이는 것입니다. '
  '마음의 눈은 보이지 않는 세계를 봅니다. 하나님의 약속, 은혜의 가능성, 미래의 소망 — '
  '이것이 마음의 눈으로 볼 수 있는 것입니다.')
blank()

p = P('성경에 이것을 잘 보여주는 이야기가 있습니다. '
  '아람 나라 군대가 엘리사를 잡으러 밤에 와서 도시를 완전히 둘러쌌습니다. '
  '아침에 이것을 본 엘리사의 사환(하인)은 기겁했습니다. "큰일났습니다! 어떡합니까?" '
  '그때 엘리사는 담담하게 말했습니다. "두려워하지 말라. 우리와 함께한 자가 '
  '저들과 함께한 자보다 많으니라." 그리고 기도했습니다. '
  '"여호와여 원컨대 저 청년의 눈을 열어서 보게 하옵소서." '
  '그러자 사환의 눈이 열렸고, 산에 불말과 불병거가 가득한 것을 보았습니다.')
fn_mgr.add(p, '열왕기하 6:14-17; "성경말씀: 열왕기하 6:14~17" (1983.2.27) 설교 참조')
blank()

P('현실은 같았습니다. 적군이 둘러싼 것은 사실이었습니다. '
  '달라진 것은 \'보는 눈\'이었습니다. '
  '마음의 눈이 열리면, 같은 상황에서도 하나님의 손길을 발견할 수 있습니다.')

# ⑤ 부정적 생각은 전염된다
numbered(5, '부정적 생각은 전염된다')
quote_box('"너희는 이 세대를 본받지 말고 오직 마음을 새롭게 함으로 변화를 받아." (로마서 12:2)')
blank()

P('부정적인 생각은 혼자 머물지 않습니다. 마음이 어둡고 부정적인 사람은 '
  '자기만 그렇게 되는 것이 아니라, 주변 사람에게도 부정적인 말을 하고 '
  '희망 없는 말을 하여 파괴적인 분위기를 퍼뜨립니다.')
blank()

P('여럿이 모여 남의 흉을 보고 수군대면 기분이 상하고 몸이 개운치 않은 것은, '
  '그 부정적인 생각이 듣는 사람의 마음에도 전달되기 때문입니다. '
  '반대로, 긍정적인 마음을 가진 사람과 함께 있으면 용기가 생기고 적극적인 마음이 생깁니다.')
blank()

P('광야에서 열 명의 정탐꾼이 부정적인 보고를 했을 때, '
  '삼백만 이스라엘 백성 전체가 울며 원망했습니다. '
  '단 열 명의 부정적인 생각이 삼백만 명에게 전염된 것입니다. '
  '그 결과, 한 세대 전체가 약속의 땅에 들어가지 못했습니다.')

key_box('핵심 정리: 생각이 중요한 5가지 이유',
        '1. 생각이 곧 나의 실체 — 마음의 생각이 어떠하면 그 사람도 그러합니다\n'
        '2. 생각은 반드시 열매를 맺음 — 좋은 씨앗이든 나쁜 씨앗이든 열매가 나옵니다\n'
        '3. 생각이 인생의 방향을 결정 — 생각→감정→태도→행동→습관→운명\n'
        '4. 마음의 눈이 현실을 해석 — 같은 상황도 다르게 볼 수 있습니다\n'
        '5. 생각은 전염됨 — 내 생각이 가정과 공동체 전체에 영향을 미칩니다')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 2부
# ══════════════════════════════════════════════════════════
doc.add_heading('2부', level=1)
doc.add_heading('왜 생각은 저절로 부정적이 되는가', level=2)

P('생각이 중요하다는 것은 누구나 알고 있습니다. 문제는 왜 우리의 생각이 '
  '자꾸만 부정적인 쪽으로 흘러가느냐 하는 것입니다. '
  '"긍정적으로 생각하세요"라는 말은 쉽지만, 실제로 마음을 바꾸는 것은 왜 이렇게 어려울까요? '
  '방법을 이야기하기 전에, 먼저 문제의 원인을 이해해야 합니다. '
  '원인을 모르면 어떤 방법도 지속되지 않습니다.')

# ① 인간의 기본값은 부정이다
numbered(1, '인간의 기본값은 부정이다')
quote_box('"육신의 생각은 사망이요 영의 생각은 생명과 평안이니라." (로마서 8:6)')
blank()

P('컴퓨터를 처음 샀을 때 기본 설정이 있듯이, 인간에게도 \'기본 설정\'이 있습니다. '
  '안타깝게도, 타락 이후 인간의 기본 설정은 \'부정\'입니다.')
blank()

P('에덴동산에서 아담과 하와는 하나님과 완벽한 교제 속에 살았습니다. '
  '그들의 생각은 밝고 긍정적이었습니다. 그런데 마귀의 거짓말에 속아 '
  '하나님의 말씀 대신 자기 감각을 따랐고, 결정적으로 하나님께 대한 감사를 잃어버렸습니다. '
  '이것이 타락의 핵심이며, 인간의 생각이 부정적으로 기울어진 시작점입니다.')
blank()

P('정원을 떠올려 보십시오. 정원사가 아무것도 하지 않으면 어떻게 됩니까? '
  '꽃이 저절로 피지 않습니다. 잡초가 무성해집니다. '
  '우리의 마음도 마찬가지입니다. 아무것도 하지 않으면 생각은 자동으로 '
  '어두운 쪽으로 흘러갑니다. 이것이 타락 이후 인간 마음의 현실입니다.')

# ② 부정적 생각의 배후에는 영적 원인이 있다
numbered(2, '부정적 생각의 배후에는 영적 원인이 있다')
P('성경은 부정적 생각을 단순한 심리 현상으로 보지 않습니다. '
  '에덴동산에서 뱀은 하와의 몸을 위협한 것이 아니라, '
  '"너도 하나님처럼 될 수 있다"는 생각을 심었습니다. '
  '가룟 유다의 배반도 마귀가 그 마음에 생각을 넣는 것으로 시작되었습니다(요 13:2). '
  '부정적 생각에는 영적 배후가 있습니다.')
blank()

p = P('조용기 목사는 이것을 누에의 비유로 설명합니다. '
  '"뽕잎을 먹고 자란 누에가 고치를 짓는데, 파리가 누에 위에 알을 까버리면 '
  '고치에서 나비가 아니라 구더기가 나옵니다. 인간의 생각도 하나님처럼 위대한 나비가 나와야 하지만, '
  '마귀가 부정적 사고의 알을 까놓으면 구더기만 나옵니다." '
  '이 비유는 우리에게 경고합니다 — 부정적 생각은 단순한 기분 문제가 아니라 '
  '영적 전쟁의 문제라는 것입니다.')
fn_mgr.add(p, '"마음의 청소" (1982.7.4) 설교에서 인용')

# ③ 방치하면 반드시 확장된다
numbered(3, '방치하면 반드시 확장된다')
P('부정적 생각은 작은 씨앗처럼 시작하지만, 방치하면 꼬리에 꼬리를 물고 자랍니다.')
blank()

P('성경에 가인의 이야기가 있습니다. 가인은 자신의 제물이 받아들여지지 않자 '
  '\'분노\'라는 작은 생각의 씨앗이 마음에 들어왔습니다. '
  '하나님이 직접 경고하셨습니다. "죄가 너를 원하나 너는 죄를 다스릴지니라"(창 4:7). '
  '그러나 가인은 그 생각을 방치했습니다. 분노는 시기가 되고, 시기는 미움이 되고, '
  '미움은 결국 살인으로 이어졌습니다. 처음에는 작은 불만이었을 뿐인데, '
  '방치한 결과 돌이킬 수 없는 비극이 되었습니다.')
blank()

P('우리의 일상에서도 마찬가지입니다. 아침에 작은 짜증 하나를 방치하면, '
  '점심때는 동료에게 짜증을 내고, 저녁에는 가족에게 화를 내고, '
  '밤에는 \'나는 왜 이 모양일까\' 하는 자기 비하로 끝나는 경우가 있습니다. '
  '작은 부정적 생각을 방치하는 순간 그것은 자라고 퍼져서 삶 전체를 지배합니다.')

# ✦ 기초: 십자가 위에서 생각을 바꾼다
foundation('기초: 십자가 위에서 생각을 바꾼다')
P('본격적인 방법을 살피기 전에, 반드시 짚어야 할 기초가 있습니다.', bold=True)
blank()

P('세상에도 \'긍정적 사고\'를 가르치는 책과 강연이 많습니다. 그런데 기독교의 생각 전환은 '
  '세상의 긍정적 사고와 근본적으로 다릅니다. 세상의 긍정적 사고는 "잘 될 거야"라는 '
  '막연한 낙관주의입니다. 그러나 기독교의 생각 전환은 역사적 사건인 십자가에 근거한 확신입니다.')
blank()

P('예수님이 십자가에서 가시관을 쓰시며 머리에서 피를 흘리셨습니다. '
  '왜 하필 머리에서 피를 흘리셨을까요? 가시관은 저주의 상징입니다. '
  '예수님이 머리(생각의 자리)에 저주의 가시관을 쓰신 것은, '
  '우리의 저주받은 생각을 대신 지시기 위함입니다. '
  '십자가의 보혈로 우리의 부정적이고 저주받은 생각이 씻겨나갈 수 있습니다.')
blank()

P('이것이 기초입니다. 이 기초 없이 아무리 긍정적으로 생각하려 해도 '
  '뿌리가 없는 나무처럼 금방 쓰러집니다. 십자가 위에서 시작할 때 '
  '비로소 생각의 전환이 진정한 힘을 얻습니다.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 3부
# ══════════════════════════════════════════════════════════
doc.add_heading('3부', level=1)
doc.add_heading('생각을 바꾸는 네 가지 실천 원리', level=2)

P('이제 구체적인 방법을 살펴보겠습니다. 설교가 제시하는 방법들은 서로 독립적이지 않습니다. '
  '네 원리는 하나의 흐름을 이룹니다.')

key_box('네 원리의 흐름',
        '원리 1. 비워라 — 잘못된 생각을 먼저 꺼냅니다 (마음의 청소)\n'
        '원리 2. 심어라 — 하나님의 말씀을 채웁니다 (생각의 프로그래밍)\n'
        '원리 3. 채워라 — 감사와 기쁨으로 체질을 바꿉니다 (마음의 건강)\n'
        '원리 4. 입어라 — 새로운 자아상을 받아들입니다 (정체성의 전환)')

divider()

# ── 원리 1 ──
P('원리 1', bold=True, size=15, color=C_NAVY, space_before=16, space_after=2)
P('마음을 비워라 — 부정적 생각의 청소', bold=True, size=13, color=C_TEAL, space_after=10)

p = P('마음의 청소는 모든 변화의 첫 단계입니다. 방에 새 가구를 들이려면 먼저 낡은 가구를 '
  '치워야 하는 것처럼, 좋은 생각을 심으려면 먼저 나쁜 생각을 비워야 합니다.')
fn_mgr.add(p, '"마음의 청소" (1982.7.4), "삶을 실패케 하는 마음가짐" (1982.1.24) 설교 참조')

sub('무엇을 비울 것인가')
P('성경은 우리 마음에서 비워내야 할 것들을 분명히 말합니다:')

P_multi([
    ('① 염려와 근심', {'bold': True, 'color': C_NAVY}),
], space_before=8)
P('"아무 것도 염려하지 말고"(빌립보서 4:6). '
  '염려는 아직 일어나지 않은 일에 대한 부정적 상상입니다. '
  '내일 무슨 일이 일어날까, 돈이 부족하면 어떡하지, 건강이 나빠지면... '
  '이런 염려를 해서 내일이 달라집니까? 달라지지 않습니다. '
  '오히려 오늘의 힘을 빼앗습니다. 예수님은 분명히 말씀하셨습니다. '
  '"내일 일을 위하여 염려하지 말라. 내일 일은 내일이 염려할 것이요, '
  '한 날의 괴로움은 그 날에 족하니라"(마 6:34).')

P_multi([
    ('② 원망과 불평', {'bold': True, 'color': C_NAVY}),
], space_before=8)
P('이스라엘 백성이 광야에서 가장 많이 한 것이 원망이었습니다. '
  '하나님이 홍해를 가르시고, 하늘에서 만나를 내려 주시고, 반석에서 물을 내셨는데도 '
  '그들은 끊임없이 원망했습니다. "차라리 애굽에서 고기 가마 곁에 앉았더라면..." '
  '원망은 하나님의 선하심을 부정하는 것이며, 마음의 독이 됩니다.')

P_multi([
    ('③ 미움과 분노', {'bold': True, 'color': C_NAVY}),
], space_before=8)
P('남을 미워하면 상대방보다 자신이 먼저 상처를 받습니다. '
  '미움은 내 마음에 독을 품고 있는 것과 같습니다. '
  '"원수를 사랑하라"는 예수님의 말씀은 상대를 위한 것만이 아니라, '
  '나 자신을 미움의 독에서 건져내기 위한 것이기도 합니다.')

P_multi([
    ('④ 열등감과 패배의식', {'bold': True, 'color': C_NAVY}),
], space_before=8)
P('"나는 안된다", "나는 부족하다", "나는 남들보다 못하다" — 이런 생각은 사실이 아닙니다. '
  '이것은 세상이 씌워놓은 거짓 안경입니다. 하나님 앞에서 당신은 '
  '"보배롭고 존귀한"(이사야 43:4) 존재입니다.')

sub('어떻게 비울 것인가 — 세 단계')

P_multi([
    ('첫째, 인식하기.', {'bold': True, 'color': C_NAVY}),
    (' 지금 내 마음에 어떤 생각이 들어 있는지 점검합니다. ', {}),
], space_before=8)
P('우리는 보통 자기 마음에 어떤 생각이 들어 있는지 의식하지 못합니다. '
  '매일 아침 5분만 조용히 앉아서 "지금 내 마음을 가장 많이 차지하고 있는 생각이 무엇인가?" '
  '물어보십시오. 마음의 청소는 자각에서 시작됩니다.')

P_multi([
    ('둘째, 거부하기.', {'bold': True, 'color': C_NAVY}),
    (' 부정적 생각이 떠오를 때, "이것은 하나님의 생각이 아니다"라고 분별하고 거부합니다.', {}),
], space_before=8)
P('생각이 떠오르는 것 자체는 죄가 아닙니다. 새가 머리 위를 날아가는 것을 막을 수는 없지만, '
  '머리 위에 둥지를 틀게 할 필요는 없습니다. 부정적 생각이 떠오르면 '
  '"예수의 이름으로 이 생각을 거부합니다"라고 선언하십시오.')

P_multi([
    ('셋째, 기도로 내려놓기.', {'bold': True, 'color': C_NAVY}),
], space_before=8)
P('"너희 구할 것을 감사함으로 하나님께 아뢰라. '
  '그리하면 모든 지각에 뛰어난 하나님의 평강이 그리스도 예수 안에서 '
  '너희 마음과 생각을 지키시리라"(빌립보서 4:6-7). '
  '염려를 기도로 바꾸는 것이 가장 확실한 비움입니다. '
  '기도는 내 짐을 하나님께 올려드리는 것입니다.')

tip_box('오늘 하루, 내 마음에 반복적으로 떠오르는 부정적 생각 하나를 찾아보십시오. '
        '그것을 종이에 적고, "주님, 이 생각을 내려놓겠습니다"라고 기도하십시오. '
        '적은 종이를 찢어버리는 것도 좋은 상징적 행동입니다.')

divider()

# ── 원리 2 ──
P('원리 2', bold=True, size=15, color=C_NAVY, space_before=16, space_after=2)
P('말씀을 심어라 — 하나님의 생각으로 프로그래밍', bold=True, size=13, color=C_TEAL, space_after=10)

p = P('마음을 비우기만 하면 오히려 더 위험합니다. 예수님은 빈집의 비유를 들어 경고하셨습니다.')
fn_mgr.add(p, '"생각을 바꿔라" (1984.11.11), "네 마음을 지키라" (1986.1.19), "마음을 새롭게 함으로 변화를 받으라" (1998.2.1) 설교 참조')

term_box(
    '빈집의 비유 (마태복음 12:43-45)',
    '더러운 귀신이 한 사람에게서 나갔다가, 돌아와 보니 그 집이 비어 있고 '
    '깨끗하게 치워져 있었습니다. 그러자 자기보다 더 악한 귀신 일곱을 데리고 돌아와 '
    '그곳에 들어가 살았습니다. 그 사람의 나중 형편이 전보다 더 나빠졌습니다.\n\n'
    '교훈: 나쁜 생각을 비우기만 하고 좋은 것으로 채우지 않으면, '
    '더 나쁜 생각이 들어옵니다. 반드시 하나님의 말씀으로 채워야 합니다.'
)
blank()

sub('빛으로 어둠을 밀어내는 원리')
P('부정적인 생각을 없애는 방법은 그 생각과 직접 싸우는 것이 아닙니다. '
  '어둠은 실체가 없습니다. 캄캄한 방에서 어둠을 주먹으로 때려본들 어둠이 없어지지 않습니다. '
  '그러나 촛불 하나만 켜면 어둠은 사라집니다. '
  '부정적인 생각도 마찬가지입니다. 그것과 싸우려 하면 도리어 더 강해집니다. '
  '그러나 생각이 하나님의 말씀으로 가득 차면 부정적인 생각이 들어올 틈이 사라집니다.')

sub('말씀으로 생각 채우기 — 구체적 실천')
P('말씀을 단순히 "읽는 것"으로 그쳐서는 안 됩니다. '
  '필요한 때에 무기로 삼아 선포할 수 있도록 마음에 새겨야 합니다.')

P_multi([
    ('① 말씀 암송', {'bold': True, 'color': C_NAVY}),
], space_before=8)
P('반복해서 외우고 또 외워, 성령의 검을 손에 쥐십시오. '
  '예수님도 광야에서 마귀의 시험을 받으실 때, 매번 "기록되었으되..."라고 말씀으로 대적하셨습니다(마 4장). '
  '예수님도 말씀을 암송하고 계셨기에 시험의 순간에 즉시 사용하실 수 있었습니다. '
  '내가 약한 영역의 말씀부터 시작하면 효과적입니다:')
bullet('염려가 많다면: "아무 것도 염려하지 말고..." (빌립보서 4:6)')
bullet('두려움이 있다면: "두려워하지 말라 내가 너와 함께 함이라." (이사야 41:10)')
bullet('자존감이 낮다면: "너는 내 눈에 보배롭고 존귀하며 내가 너를 사랑하노라." (이사야 43:4)')

P_multi([
    ('② 말씀 묵상', {'bold': True, 'color': C_NAVY}),
], space_before=8)
P('"복 있는 사람은 오직 여호와의 율법을 즐거워하여 주야로 묵상하는도다"(시편 1:2). '
  '묵상은 말씀을 천천히 씹어서 단맛이 나오게 하는 것입니다. '
  '한 구절을 읽고, "이 말씀이 나에게 무엇을 의미하는가?", '
  '"오늘 내 상황에 이것을 어떻게 적용할 수 있는가?" 생각해 보는 것입니다.')

P_multi([
    ('③ 말씀에 내 이름 넣어 읽기', {'bold': True, 'color': C_NAVY}),
], space_before=8)
P('"내게 능력 주시는 자 안에서 내가 모든 것을 할 수 있느니라"(빌 4:13)를 읽을 때, '
  '자기 이름을 넣어 읽어 보십시오. "OO에게 능력 주시는 자 안에서 OO가 모든 것을 할 수 있느니라." '
  '추상적 진리가 나의 현실로 내려옵니다.')

sub('오중복음 묵상 — 생각에 채울 구체적 내용')
P('무엇을 묵상해야 하는가? 조용기 목사는 매우 구체적인 답을 줍니다. 바로 오중복음입니다. '
  '"나는 성도들에게 늘 오중복음을 생각하라고 가르칩니다. '
  '오중복음의 내용으로 마음을 채우면 부정적인 생각이 들어올 틈이 없습니다."')

term_box(
    '오중복음(五重福音)이란?',
    '조용기 목사가 평생 가르친 다섯 가지 핵심 복음의 내용입니다. '
    '이 다섯 가지를 생각에 채우는 것이 묵상의 핵심입니다.\n\n'
    '① 중생(重生)의 복음 — 거듭남\n'
    '  예수 그리스도를 믿으면 죄 사함을 받고 완전히 새 사람이 됩니다.\n'
    '  "누구든지 그리스도 안에 있으면 새로운 피조물이라"(고후 5:17)\n'
    '  → 묵상: "나는 이미 용서받았다. 나는 새로운 피조물이다."\n\n'
    '② 성령충만의 복음 — 성령님의 도우심\n'
    '  성령님이 우리 안에 오셔서 능력과 지혜와 위로를 주십니다.\n'
    '  "오직 성령이 너희에게 임하시면 너희가 권능을 받고"(행 1:8)\n'
    '  → 묵상: "나 혼자 싸우는 것이 아니다. 성령님이 나와 함께 계신다."\n\n'
    '③ 신유(神癒)의 복음 — 치유\n'
    '  예수님의 채찍 맞으신 상처로 우리의 영혼과 몸이 나음을 입었습니다.\n'
    '  "그가 채찍에 맞음으로 너희는 나음을 얻었나니"(벧전 2:24)\n'
    '  → 묵상: "하나님은 나를 치유하시는 분이시다."\n\n'
    '④ 축복의 복음 — 삼중축복\n'
    '  하나님은 영혼의 잘됨, 건강, 물질적 필요까지 돌보십니다.\n'
    '  "네 영혼이 잘됨 같이 네가 범사에 잘되고 강건하기를"(요삼 1:2)\n'
    '  → 묵상: "하나님은 나의 모든 필요를 아시고 채워주신다."\n\n'
    '⑤ 재림(再臨)의 복음 — 소망\n'
    '  예수님이 다시 오시며, 부활과 영원한 천국이 우리를 기다립니다.\n'
    '  "이 예수는 너희 가운데서 하늘로 올려지심을 본 그대로 오시리라"(행 1:11)\n'
    '  → 묵상: "이 세상이 끝이 아니다. 영원한 소망이 있다."'
)
blank()

P('이 다섯 가지 진리를 매일 묵상하면, 마음이 어둠으로 채워질 틈이 없습니다. '
  '특히 힘든 상황에서 "나는 이미 용서받았고(중생), 성령님이 함께하시며(성령충만), '
  '치유하시는 하나님이 계시고(신유), 모든 필요를 채워주시며(축복), '
  '영원한 소망이 있다(재림)"고 묵상하면 어떤 절망도 버틸 수 있습니다.')

tip_box('이번 주에 오중복음 중 하나를 골라 관련 말씀 한 구절을 암송하십시오.\n'
        '예: 자존감이 약하다면 → 중생의 복음 → "고후 5:17" 암송\n'
        '   불안하다면 → 성령충만의 복음 → "행 1:8" 암송\n'
        '아침에 일어나서, 밤에 잠들기 전에 그 말씀을 소리 내어 읽고 묵상하십시오.')

divider()

# ── 원리 3 ──
P('원리 3', bold=True, size=15, color=C_NAVY, space_before=16, space_after=2)
P('감사와 기쁨으로 채워라 — 마음의 체질을 바꾸는 훈련', bold=True, size=13, color=C_TEAL, space_after=10)

p = P('말씀으로 생각을 채웠다면, 이제 그 생각을 일상에서 유지하고 강화해야 합니다. '
  '가장 강력한 도구가 바로 감사와 기쁨입니다.')
fn_mgr.add(p, '"축복을 가져오는 생활태도" (1996.10.6), "감사보은의 생활" (1982.11.21), "기쁨의 샘" (1985.9.1), "마음의 건강" (1985.9.22) 설교 참조')

sub('감사는 감정이 아니라 결단이다')
quote_box('"항상 기뻐하라. 쉬지 말고 기도하라. 범사에 감사하라.\n이것이 그리스도 예수 안에서 너희를 향하신 하나님의 뜻이니라." (데살로니가전서 5:16-18)')
blank()

P('성경은 \'기분이 좋을 때\' 감사하라고 하지 않습니다. \'범사에\' — 모든 일에 — 감사하라고 합니다. '
  '이것은 감정의 문제가 아니라 의지의 결단입니다.')
blank()

P('감사가 어려운 이유는, 우리가 \'없는 것\'에 집중하기 때문입니다. '
  '"왜 월급이 이것밖에 안 되지?", "왜 내 집은 이렇게 좁지?", '
  '"왜 내 남편(아내)은 저 사람 남편(아내)과 다르지?" '
  '그러나 감사하는 사람은 \'있는 것\'에 집중합니다. '
  '"일할 수 있는 건강이 있으니 감사합니다", "비바람을 피할 집이 있으니 감사합니다", '
  '"함께 살아가는 가족이 있으니 감사합니다." '
  '같은 현실인데, 보는 방향만 바뀌면 원망이 감사로 변합니다.')

sub('기쁨은 마음의 양약이다')
quote_box('"마음의 기쁨은 양약이라도 심령의 상함은 뼈의 썩음이라." (잠언 17:22)')
blank()

p = P('사도 바울이 빌립보서를 쓸 때, 그는 로마의 감옥에 있었습니다. '
  '그곳은 땅 속 지하 감옥으로, 캄캄하고 습기 차고 벼룩과 빈대가 득실거리는 곳이었습니다. '
  '언제 풀려날지 기약도 없었습니다. 그런데 바울은 이 감옥에서 이렇게 썼습니다. '
  '"주 안에서 항상 기뻐하라. 내가 다시 말하노니 기뻐하라"(빌 4:4). '
  '만일 바울이 그 칠흑 같은 어둠 속에서 우울하고 절망했더라면, '
  '얼마 있지 않아 정신적으로 무너졌을 것입니다. '
  '그러나 넘치는 기쁨이 그의 마음을 지켜주었습니다.')
fn_mgr.add(p, '"마음의 건강" (1985.9.22) 설교에서 바울의 상황을 상세히 다룸')
blank()

P('기쁨은 환경에서 오지 않습니다. 환경이 좋아야 기뻐할 수 있다면, '
  '이 세상에서 항상 기뻐할 수 있는 사람은 아무도 없습니다. '
  '주 안에서 오는 기쁨은 환경과 상관없이 마음을 지키는 보호막입니다.')

sub('감사와 기쁨의 구체적 실천')

P_multi([
    ('① 감사 일기', {'bold': True, 'color': C_NAVY}),
], space_before=8)
P('매일 잠들기 전에 감사한 것 3가지를 적어봅니다. '
  '처음에는 억지로 느껴질 수 있습니다. "감사할 게 뭐가 있지?" 싶을 수도 있습니다. '
  '그러나 한 달만 지속하면 놀라운 변화가 일어납니다. '
  '감사의 눈이 열리기 시작하면, 전에는 보이지 않던 은혜가 보이기 시작합니다.')

P_multi([
    ('② 찬양의 시간', {'bold': True, 'color': C_NAVY}),
], space_before=8)
P('어려울 때일수록 찬양하십시오. 마음이 무거울 때 찬송가를 부르면 '
  '마음에 빛이 들어옵니다. 찬양은 마음의 분위기를 바꾸는 강력한 도구입니다. '
  '사울 왕이 악한 영에 시달릴 때, 다윗의 수금 연주(찬양)가 그를 치유했습니다(삼상 16:23).')

P_multi([
    ('③ 불평 대신 감사로 바꾸기', {'bold': True, 'color': C_NAVY}),
], space_before=8)
P('연습이 필요합니다. 불평이 나오려 할 때, 의식적으로 감사로 바꿔보십시오:')
bullet('"날씨가 너무 덥다" → "건강하게 여름을 보내게 해주셔서 감사합니다"')
bullet('"일이 너무 많다" → "일할 수 있는 능력을 주셔서 감사합니다"')
bullet('"아이가 말을 안 들어" → "이 아이를 주신 것이 감사합니다"')
P('이것은 현실을 부정하는 것이 아닙니다. 같은 현실을 하나님의 관점에서 다시 보는 것입니다.')

tip_box('오늘부터 잠들기 전, 감사한 것 3가지를 떠올리며 하나님께 감사 기도를 드려보십시오.\n'
        '작은 것부터 시작하세요. "오늘 밥을 먹을 수 있었습니다" — 이것도 감사입니다.\n'
        '한 달 뒤, 당신의 마음이 얼마나 달라졌는지 놀라게 될 것입니다.')

divider()

# ── 원리 4 ──
P('원리 4', bold=True, size=15, color=C_NAVY, space_before=16, space_after=2)
P('새 자아상을 입어라 — 하나님의 눈으로 나를 보기', bold=True, size=13, color=C_TEAL, space_after=10)

p = P('생각을 바꾸는 가장 근본적인 작업은 \'나는 누구인가\'에 대한 답을 바꾸는 것입니다. '
  '391편의 "생각" 설교 가운데 가장 높은 점수를 받은 상위 설교들이 모두 "자화상"에 관한 설교입니다. '
  '그만큼 이 주제는 생각의 영역에서 핵심 중의 핵심입니다.')
fn_mgr.add(p, '"보혈로 그린 자화상" (1992.3.1), "삶의 성공과 실패를 가져오는 자화상" (2001.7.1), "나는 나의 자화상을 본다" (2015.3.15), "자화상" (2017.10.29) 설교 참조')

sub('자화상이란 무엇인가')
term_box(
    '자화상(自畵像)이란?',
    '자화상은 "내가 나를 보는 모습"입니다. 화가가 거울을 보고 자기 모습을 그리듯이, '
    '우리 마음속에도 나 자신에 대한 그림이 있습니다. '
    '이 그림이 밝으면 밝은 삶을 살고, 어두우면 어두운 삶을 삽니다.\n\n'
    '색안경의 비유: 파란 안경을 끼면 세상이 파랗게 보이듯, '
    '마음속에 어떤 자화상을 품고 있느냐에 따라 같은 세상이 '
    '절망으로도, 희망으로도 보입니다.'
)
blank()

sub('세상의 자화상 vs 하나님의 자화상')
P('세상은 끊임없이 우리에게 자화상을 그려줍니다:')
bullet('성적, 학벌, 외모, 재산으로 사람의 가치를 매깁니다.')
bullet('실패하면 \'실패자\'라는 딱지를 붙입니다.')
bullet('남과 비교하고, 경쟁하게 하고, 부족함을 느끼게 합니다.')
bullet('"넌 안돼", "넌 부족해", "넌 자격이 없어" — 이런 말이 반복되면 그것이 내 자화상이 됩니다.')
blank()

P('그러나 하나님은 전혀 다른 자화상을 보여주십니다:')
quote_box('"너는 내 눈에 보배롭고 존귀하며 내가 너를 사랑하노라." (이사야 43:4)')
blank()
bullet('"너희는 택하신 족속이요 왕 같은 제사장들이요 거룩한 나라요 그의 소유가 된 백성이니." (벧전 2:9)')
bullet('"누구든지 그리스도 안에 있으면 새로운 피조물이라. 이전 것은 지나갔으니 보라 새 것이 되었도다." (고후 5:17)')
blank()

p = P('조용기 목사는 자화상의 중요성을 이렇게 강조합니다. '
  '"인간은 모두다 자기가 자신을 보고 느끼는 모습을 가슴에 품고, '
  '그 모습을 쫓아 말하고 생각하고 행동하며 삽니다. '
  '자화상이 열등하면 열등의식에 잡히고, 자화상이 성공적이면 당당한 사람으로 행동하게 됩니다."')
fn_mgr.add(p, '"보혈로 그린 자화상" (1992.3.1) 설교에서 인용')

sub('다윗의 자화상 — 목동에서 왕으로')
P('어린 다윗은 막내아들로, 형들이 전쟁터에 갈 때 혼자 양을 치는 목동이었습니다. '
  '아버지 이새도 그를 대수롭지 않게 여겼습니다. 사무엘 선지자가 찾아왔을 때 '
  '아버지는 다윗을 부르지도 않았습니다. 세상의 눈으로 다윗은 \'보잘것없는 막내\'였습니다.')
blank()

P('그러나 다윗의 마음속 자화상은 달랐습니다. 골리앗 앞에 섰을 때, '
  '온 이스라엘 군대가 두려워 떨고 있었지만 다윗은 이렇게 말했습니다. '
  '"너는 칼과 창과 단창으로 내게 나아오거니와 나는 만군의 여호와의 이름으로 네게 나아가노라"(삼상 17:45). '
  '다윗은 자신을 \'막내 목동\'으로 보지 않았습니다. \'하나님의 전사\'로 보았습니다. '
  '이 자화상이 그를 왕으로 이끌었습니다.')

sub('새 자아상을 세우는 구체적 방법')

P_multi([
    ('① 하나님이 나에 대해 말씀하신 것을 찾기', {'bold': True, 'color': C_NAVY}),
], space_before=8)
P('성경에서 \'하나님의 자녀\'가 누구인지 찾아 읽으십시오. '
  '에베소서 1장, 로마서 8장, 베드로전서 2장에 놀라운 선언들이 있습니다. '
  '그것이 세상이 당신에게 붙인 딱지가 아니라 당신의 참 모습입니다.')

P_multi([
    ('② 매일 고백하기', {'bold': True, 'color': C_NAVY}),
], space_before=8)
P('아침마다 거울 앞에서, 또는 조용한 기도 시간에 이렇게 고백해 보십시오:')
bullet('"나는 하나님의 자녀입니다."')
bullet('"나는 그리스도 안에서 새로운 피조물입니다."')
bullet('"나는 보배롭고 존귀한 존재입니다."')
bullet('"나는 용서받았고, 사랑받고 있으며, 하나님의 계획 안에 있습니다."')
P('이것은 자기 최면이 아닙니다. 하나님의 말씀에 근거한 진실의 고백입니다. '
  '처음에는 어색하게 느껴질 수 있습니다. 그러나 이 고백이 마음 깊은 곳에 뿌리내리면 '
  '삶의 방향이 달라지기 시작합니다.')

P_multi([
    ('③ 과거의 실패에서 벗어나기', {'bold': True, 'color': C_NAVY}),
], space_before=8)
P('많은 사람이 과거의 실패에 묶여 삽니다. "나는 그때 실패했으니까...", '
  '"나는 그때 잘못했으니까..." 그러나 성경은 선언합니다. '
  '"이전 것은 지나갔으니 보라 새 것이 되었도다"(고후 5:17). '
  '하나님은 과거가 아니라 미래를 보시는 분입니다. '
  '실패의 기억이 떠오를 때마다, 이 말씀을 묵상하십시오.')

tip_box('거울 앞에서 "나는 하나님이 사랑하시는 자녀입니다"라고 고백해 보십시오.\n'
        '처음에는 어색하지만, 이 고백이 자화상을 바꾸는 출발점입니다.\n'
        '매일 아침 1분이면 됩니다. 한 달만 해보십시오.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 나가며
# ══════════════════════════════════════════════════════════
doc.add_heading('나가며', level=1)
doc.add_heading('생각이 바뀌면 삶이 바뀐다', level=2)

quote_box('"우리가 낙심하지 아니하노니 우리의 겉사람은 낡아지나\n우리의 속사람은 날로 새로워지도다." (고린도후서 4:16)')
blank()

P('이 네 가지 원리는 한 번의 결심이 아니라 매일의 훈련입니다. '
  '농부가 매일 밭을 가꾸듯이, 우리도 매일 마음의 밭을 가꿔야 합니다. '
  '어떤 날은 잘 되고, 어떤 날은 다시 부정적인 생각이 밀려올 것입니다. '
  '괜찮습니다. 넘어져도 다시 일어나면 됩니다. 중요한 것은 포기하지 않는 것입니다.')
blank()

P('생각을 바꾸는 것은 하루아침에 되지 않습니다. '
  '그러나 오늘 한 걸음을 시작하면, 내일은 어제보다 한 걸음 앞서 있을 것입니다. '
  '하나님은 우리의 한 걸음 한 걸음을 보시며 함께 걸어주십니다.')
blank()

quote_box('"무릇 진실하며 무릇 존귀하며 무릇 옳으며 무릇 정결하며\n'
          '무릇 사랑 받을 만하며 무릇 칭찬 받을 만하며\n'
          '무슨 덕이 있든지 무슨 기림이 있든지 이것들을 생각하라." (빌립보서 4:8)')

blank(8)

# ── 정리표 ──
P('네 가지 원리 한눈에 보기', bold=True, size=12, color=C_NAVY, space_before=12)
blank()

table = doc.add_table(rows=5, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
# 헤더
headers = ['원리', '핵심 행동', '방향', '핵심 말씀']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = C_WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 배경색
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1B4F72')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

data = [
    ['1. 비워라', '염려·원망·미움을\n내려놓기', '마음의 청소', '빌 4:6'],
    ['2. 심어라', '말씀 암송·묵상\n오중복음 채우기', '생각의\n프로그래밍', '시 1:2'],
    ['3. 채워라', '감사·기쁨으로\n마음 체질 바꾸기', '마음의 건강', '살전 5:16-18'],
    ['4. 입어라', '보혈의 자화상을\n받아들이기', '정체성의 전환', '고후 5:17'],
]
colors = ['EBF5FB', 'FEF9E7', 'EAFAF1', 'F5EEF8']
for r, (row_data, bg) in enumerate(zip(data, colors)):
    for c, val in enumerate(row_data):
        cell = table.rows[r+1].cells[c]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if c == 0:
            run.bold = True
            run.font.color.rgb = C_NAVY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), bg)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

blank(12)

# ── 기도문 ──
P('함께 드리는 기도', bold=True, size=12, color=C_TEAL, space_before=16)
blank()

P('"하나님 아버지, 제 생각을 주님의 생각으로 바꿔주십시오. '
  '제 마음에 쌓인 염려와 두려움과 열등감을 십자가 앞에 내려놓습니다. '
  '말씀으로 제 생각을 채워주시고, 감사와 기쁨으로 제 마음의 체질을 바꿔주십시오. '
  '세상이 그려놓은 거짓 자화상을 벗기시고, '
  '보혈로 그려주신 새 자화상을 받아들이게 하여 주십시오. '
  '오늘 한 걸음, 주님과 함께 시작하겠습니다. '
  '예수님의 이름으로 기도합니다. 아멘."', italic=True, color=C_SUBTLE)


# ═══════════════════════════════════════════════════════════
# 저장
# ═══════════════════════════════════════════════════════════
output_path = 'output/4차원영성_생각_아티클.docx'
doc.save(output_path)
print(f'생성 완료: {output_path}')
print(f'총 단락 수: {len(doc.paragraphs)}')
print(f'총 표 수: {len(doc.tables)}')
print(f'총 각주 수: {fn_mgr.next_id - 1}')
