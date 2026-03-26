# -*- coding: utf-8 -*-
"""
4차원영성 생각 아티클 → .docx 변환 스크립트
생각의_통합본.docx 형식을 참조하여 생성
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── 스타일 설정 ──────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.5

for level in [1, 2]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '맑은 고딕'
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    if level == 1:
        hs.font.size = Pt(18)
        hs.font.bold = True
    else:
        hs.font.size = Pt(14)
        hs.font.bold = True

# ── 헬퍼 함수 ──────────────────────────────────────────
def add_normal(text, bold=False, italic=False, size=None, color=None, align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_quote(text, space_before=6):
    """성경 인용구 (이탤릭, 들여쓰기)"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def add_ref(text):
    """설교 참조 (작은 회색 글씨)"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_sub(text):
    """▌ 소제목"""
    p = doc.add_paragraph()
    run = p.add_run(f'▌ {text}')
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(text):
    """• 불릿 포인트"""
    p = doc.add_paragraph()
    run = p.add_run(f'•  {text}')
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_numbered(num, text):
    """① ② ③ 넘버링"""
    circled = ['①','②','③','④','⑤','⑥','⑦','⑧','⑨','⑩']
    p = doc.add_paragraph()
    run = p.add_run(f'{circled[num-1]} {text}')
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_foundation(text):
    """✦ 기초 섹션"""
    p = doc.add_paragraph()
    run = p.add_run(f'✦  {text}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    return p

def blank():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

# ══════════════════════════════════════════════════════════
# 표지 정보
# ══════════════════════════════════════════════════════════
add_normal('4차원의 영성 · 생각(Thinking)', size=10, color=RGBColor(0x66,0x66,0x66), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_normal("'생각'의 능력", size=9, color=RGBColor(0x88,0x88,0x88), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_normal('조용기 목사 설교 분석', size=9, color=RGBColor(0x88,0x88,0x88), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_normal('— 설교 원문(1981~2020년, 391편) 기반 —', size=9, color=RGBColor(0x88,0x88,0x88), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

# ══════════════════════════════════════════════════════════
# 들어가며
# ══════════════════════════════════════════════════════════
doc.add_heading('들어가며', level=1)
doc.add_heading('생각이 인생을 만든다', level=2)

add_quote('"지킬만한 것보다 더욱 네 마음을 지키라. 생명의 근원이 이에서 남이니라." (잠언 4:23)')
blank()

add_normal(
    '성경이 재산도, 건강도, 인간관계도 아닌 \'마음\'을 가장 먼저 지키라고 말씀하는 이유는 무엇일까요? '
    '그것은 우리의 생각이 단순히 머릿속에서 지나가는 관념이 아니라, 삶의 방향과 결과를 실질적으로 결정하는 '
    '\'생명의 근원\'이기 때문입니다.'
)
blank()

add_normal(
    '이 명제는 단순한 긍정적 사고 구호가 아닙니다. 4차원의 영성이라는 신학적 틀에서 나온 원리입니다. '
    '조용기 목사에 따르면, 보이지 않는 4차원(영적 세계)이 보이는 3차원(물질 세계)을 지배합니다. '
    '1차원은 2차원 안에, 2차원은 3차원 안에, 3차원은 4차원 안에 있습니다. '
    '그리고 이 4차원을 구성하는 네 가지 요소가 있습니다 — 생각, 꿈, 믿음, 말입니다. '
    '그 가운데 생각이 첫 번째입니다. 생각이 바뀌지 않으면 꿈도, 믿음도, 말도 바뀌지 않습니다. '
    '생각은 모든 변화의 출발점입니다.'
)
blank()

add_normal(
    '그렇다면 핵심 질문이 남습니다. 생각을 바꾸고 싶다는 바람은 누구나 품고 있습니다. '
    '그러나 어떻게 해야 생각이 실제로 바뀌는 것일까요? 의지만으로는 되지 않습니다. '
    '감정은 믿을 수 없습니다. 환경이 바뀌기를 기다리는 것은 더욱 위험합니다. '
    '이 문서는 조용기 목사의 설교 391편에서 추출한 원리를 바탕으로, '
    '생각이 왜 중요한지, 왜 부정적이 되는지, 그리고 어떻게 바꿀 수 있는지를 정리합니다.'
)

# ══════════════════════════════════════════════════════════
# 1부
# ══════════════════════════════════════════════════════════
doc.add_heading('1부', level=1)
doc.add_heading('생각은 왜 중요한가', level=2)

# ① 생각이 곧 그 사람의 실체이다
add_numbered(1, '생각이 곧 그 사람의 실체이다')
add_normal(
    '성경은 이 원리를 명확하게 선언합니다. "그 마음의 생각이 어떠하면 그 사람도 그러하니"(잠 23:7). '
    '겉으로 보이는 외모, 직업, 사회적 지위가 그 사람의 진정한 모습이 아닙니다. '
    '그 사람의 마음속에 어떤 생각이 자리잡고 있느냐가 그 사람의 진정한 모습입니다.'
)
blank()

add_normal(
    '인간이 지구의 주인이 된 것은 체력 때문이 아닙니다. 사자보다 힘이 약하고, '
    '독수리처럼 날 수도 없으며, 물고기처럼 물속에서 살 수도 없습니다. '
    '그런데도 인간이 이 지구를 정복한 것은 오직 \'생각할 수 있는 능력\' 때문입니다.'
)

# ② 생각은 씨앗처럼 반드시 열매를 맺는다
add_numbered(2, '생각은 씨앗처럼 반드시 열매를 맺는다')
add_normal(
    '예수님의 씨 뿌리는 자의 비유(마 13장)는 마음과 생각의 관계를 설명합니다. '
    '네 종류의 밭 — 길가, 자갈밭, 가시밭, 옥토 — 은 생각의 상태를 보여줍니다. '
    '매일의 생각이 곧 매일 마음의 밭에 뿌리는 씨앗입니다.'
)
blank()

add_normal(
    '"나는 못한다, 나는 안 된다, 나는 실패자다" — 이런 부정적 생각의 씨앗을 심으면 '
    '반드시 파멸의 열매를 거두게 됩니다. 반대로 "하나님이 나와 함께하신다, 나는 할 수 있다, '
    '나는 축복받았다" — 이런 생각의 씨앗은 성공과 축복의 열매를 맺습니다.'
)

add_normal(
    '생각 → 감정 → 태도 → 행동 → 습관 → 인격 → 운명. '
    '이 연쇄의 출발점이 바로 생각입니다. 같은 환경, 같은 출발점에서도 '
    '생각의 방향에 따라 한 사람은 성공하고 다른 사람은 실패합니다. '
    '이것은 운명이 아닙니다. 생각의 선택입니다.'
)

# ③ 겉사람과 속사람 — 보이지 않는 내면이 진짜 나
add_numbered(3, '겉사람과 속사람 — 보이지 않는 내면이 진짜 나')
add_normal(
    '사도 바울은 인간을 "겉사람"과 "속사람"으로 구분합니다(고후 4:16). '
    '겉사람, 즉 육체는 시간이 지남에 따라 쇠해갑니다. '
    '그러나 속사람, 즉 내면의 영과 생각은 날로 새로워질 수 있습니다. '
    '진정한 "나"는 겉으로 보이는 육체가 아니라 안에 있는 생각과 영입니다.'
)
blank()

add_normal(
    '재산을 잃어도 마음이 건강하면 다시 일어설 수 있습니다. '
    '건강을 잃어도 마음이 살아 있으면 희망이 있습니다. '
    '그러나 마음을 잃으면 — 생각이 부패하고 절망에 빠지면 — '
    '모든 것을 가지고 있어도 아무 소용이 없습니다. '
    '그래서 성경은 무엇보다 먼저 마음을 지키라고 명합니다.'
)

# ══════════════════════════════════════════════════════════
# 2부
# ══════════════════════════════════════════════════════════
doc.add_heading('2부', level=1)
doc.add_heading('왜 생각은 저절로 부정적이 되는가', level=2)

add_normal(
    '방법을 이야기하기 전에, 먼저 왜 생각이 부정적으로 흐르는지를 이해해야 합니다. '
    '문제의 구조를 모르면 어떤 방법도 지속되지 않습니다. '
    '설교는 부정적 생각의 원인을 세 가지 층위에서 진단합니다.'
)

# ① 인간의 기본값은 부정이다
add_numbered(1, '인간의 기본값은 부정이다')
add_normal(
    '인간은 타락 이후 저주받은 세상에서 태어납니다. '
    '아담과 하와는 마귀의 거짓말에 속아 하나님의 말씀 대신 자신의 생각과 감각을 따랐고, '
    '결정적으로 하나님께 대한 감사를 잃어버렸습니다. '
    '이것이 타락의 핵심이며, 생각이 부정적으로 기울어진 시작점입니다.'
)
blank()

add_normal(
    '"육신의 생각은 사망이요 영의 생각은 생명과 평안이니라"(롬 8:6). '
    '타락 이후 인간의 기본값은 부정입니다. 아무것도 하지 않으면 생각은 자동으로 '
    '어두운 쪽으로 흘러갑니다. 정원을 방치하면 잡초가 자라듯, '
    '마음을 방치하면 부정적 생각이 가득 찹니다.'
)

# ② 부정적 생각의 배후에는 마귀가 있다
add_numbered(2, '부정적 생각의 배후에는 마귀가 있다')
add_normal(
    '설교는 부정적 생각을 단순한 심리 현상으로 보지 않습니다. '
    '에덴동산에서 뱀은 하와의 몸을 위협한 것이 아니라 "너도 하나님처럼 될 수 있다"는 생각을 심었습니다. '
    '가룟 유다의 배반도 마귀가 그 마음에 생각을 넣는 것으로 시작되었습니다(요 13:2).'
)
blank()

add_normal(
    '설교는 이것을 비유로 설명합니다. "뽕잎을 먹고 자란 누에가 고치를 짓는데, '
    '파리가 누에 위에 알을 까버리면 고치에서 나비가 아니라 구더기가 나온다. '
    '인간의 생각도 하나님처럼 위대한 나비가 나와야 하지만, '
    '마귀가 부정적 사고의 알을 까놓으면 구더기만 나온다." '
    '부정적 생각에는 영적 배후가 있으므로, 영적 무기로 대적해야 합니다.'
)

# ③ 방치하면 반드시 확장된다
add_numbered(3, '방치하면 반드시 확장된다')
add_normal(
    '부정적 생각은 작은 씨앗처럼 시작하지만, 방치하면 꼬리에 꼬리를 물고 자랍니다. '
    '가인의 분노가 결국 살인으로 이어졌던 것처럼, '
    '이스라엘 백성에게 재앙이 임한 것도 그들의 생각의 결과였습니다(렘 6:19). '
    '부정적 생각은 방치하는 순간 자라고 퍼져서, 마침내 삶 전체를 지배합니다.'
)

# ✦ 기초: 십자가 위에서 생각을 바꾼다
add_foundation('기초: 십자가 위에서 생각을 바꾼다')
add_normal(
    '본격적인 방법을 살피기 전에, 반드시 짚어야 할 기초가 있습니다.'
)
blank()

add_normal(
    '세상의 긍정적 사고와 기독교의 생각 전환은 근본적으로 다릅니다. '
    '막연한 낙관주의가 아니라, 역사적 사건인 십자가에 근거한 확신입니다. '
    '예수님이 가시관을 쓰시며 머리에서 피를 흘리셨기에 우리의 저주받은 생각이 씻겨나갈 수 있습니다. '
    '십자가는 생각 전환의 기초이며, 이 기초 없이는 어떤 방법도 뿌리를 내리지 못합니다.'
)

# ══════════════════════════════════════════════════════════
# 3부
# ══════════════════════════════════════════════════════════
doc.add_heading('3부', level=1)
doc.add_heading('생각을 바꾸는 네 가지 실천 원리', level=2)

add_normal(
    '설교가 제시하는 방법들은 서로 독립적이지 않습니다. '
    '각 방법이 어떤 논리로 생각에 작용하는지, 서로 어떤 관계 속에 있는지를 이해할 때 '
    '비로소 실제로 적용할 수 있습니다. 네 원리는 하나의 흐름을 이룹니다.'
)
blank()

add_normal(
    '먼저 잘못된 생각을 비우고(원리 1), 하나님의 말씀으로 채우며(원리 2), '
    '감사와 기쁨으로 마음의 체질을 바꾸고(원리 3), '
    '새로운 자아상을 입습니다(원리 4).'
)

# ── 원리 1 ──
add_ref('관련 설교: "마음의 청소", "우리 속사람을 대적하는 원수들", "삶을 낭패케 하는 것들", "삶을 실패케 하는 마음가짐"')

add_normal('원리 1. 마음을 비워라 — 부정적 생각의 청소', bold=True, size=13, space_before=16, space_after=8)

add_sub('무엇을 비울 것인가')
add_normal(
    '생각을 바꾸려면, 먼저 잘못된 생각을 비워야 합니다. '
    '방에 새 가구를 들이려면 먼저 낡은 가구를 치워야 하는 것과 같습니다. '
    '성경은 우리 마음에서 비워내야 할 것들을 분명히 말합니다.'
)

add_bullet('염려와 근심 — "아무 것도 염려하지 말고"(빌 4:6). 염려는 아직 일어나지 않은 일에 대한 부정적 상상입니다.')
add_bullet('원망과 불평 — 원망은 하나님의 선하심을 부정하는 것이며, 마음의 독이 됩니다.')
add_bullet('미움과 분노 — 남을 미워하면 상대방보다 자신이 먼저 상처를 받습니다.')
add_bullet('열등감과 패배의식 — "나는 안된다", "나는 부족하다"는 생각은 세상이 씌워놓은 거짓 안경입니다.')

add_sub('어떻게 비울 것인가')
add_bullet('인식하기: 지금 내 마음에 어떤 생각이 들어 있는지 점검합니다. 마음의 청소는 자각에서 시작됩니다.')
add_bullet('거부하기: 부정적 생각이 떠오를 때, "이것은 하나님의 생각이 아니다"라고 분별하고 거부합니다.')
add_bullet('기도로 내려놓기: "너희 구할 것을 감사함으로 하나님께 아뢰라"(빌 4:6). 염려를 기도로 바꾸는 것이 가장 확실한 비움입니다.')

# ── 원리 2 ──
add_ref('관련 설교: "생각을 바꿔라", "영의 새로운 것으로 살아라", "네 마음을 지키라", "마음을 새롭게 함으로 변화를 받으라"')

add_normal('원리 2. 말씀을 심어라 — 하나님의 생각으로 프로그래밍', bold=True, size=13, space_before=16, space_after=8)

add_sub('빛으로 어둠을 밀어내는 원리')
add_normal(
    '부정적인 생각을 없애는 방법은 그 생각과 직접 싸우는 것이 아닙니다. '
    '어둠은 실체가 없습니다. 빛이 들어오면 어둠은 있을 자리가 없어집니다. '
    '부정적인 생각도 마찬가지입니다. 생각이 하나님의 말씀으로 가득 차면 '
    '부정적인 생각이 들어올 틈이 사라집니다.'
)

add_sub('말씀으로 생각 채우기 — 구체적 실천')
add_bullet('말씀 암송: 반복해서 외우고 또 외워, 성령의 검을 손에 쥐십시오.')
add_bullet('말씀 묵상: 복 있는 사람은 오직 여호와의 율법을 즐거워하여 주야로 묵상합니다(시 1:2). 묵상은 말씀을 씹어서 단물이 나오게 하는 것입니다.')
add_bullet('말씀에 내 이름 넣어 읽기: "내게 능력 주시는 자 안에서 내가 모든 것을 할 수 있느니라"를 읽을 때, 나의 이름을 넣어 읽으십시오. 추상적 진리가 나의 현실로 내려옵니다.')

add_sub('오중복음 묵상 — 생각에 채울 구체적 내용')
add_normal(
    '무엇을 묵상해야 하는가? 설교는 매우 구체적인 답을 줍니다. 바로 오중복음입니다. '
    '죄 사함의 은혜, 성령 충만, 치유, 축복, 부활과 천국. '
    '이 다섯 가지 복음의 내용을 생각에 채우는 것이 묵상의 핵심입니다.'
)

# ── 원리 3 ──
add_ref('관련 설교: "축복을 가져오는 생활태도", "감사보은의 생활", "감사의 제사", "기쁨의 샘", "마음의 건강", "마음의 평강"')

add_normal('원리 3. 감사와 기쁨으로 채워라 — 마음의 체질을 바꾸는 훈련', bold=True, size=13, space_before=16, space_after=8)

add_sub('감사는 감정이 아니라 결단이다')
add_quote('"항상 기뻐하라. 쉬지 말고 기도하라. 범사에 감사하라. 이것이 그리스도 예수 안에서 너희를 향하신 하나님의 뜻이니라." (살전 5:16-18)')
add_normal(
    '성경은 \'기분이 좋을 때\' 감사하라고 하지 않습니다. \'범사에\' — 모든 일에 — 감사하라고 합니다. '
    '이것은 감정의 문제가 아니라 의지의 결단입니다. '
    '감사가 어려운 이유는, 우리가 \'없는 것\'에 집중하기 때문입니다. '
    '그러나 감사하는 사람은 \'있는 것\'에 집중합니다.'
)

add_sub('기쁨은 마음의 양약이다')
add_quote('"마음의 기쁨은 양약이라도 심령의 상함은 뼈의 썩음이라." (잠언 17:22)')
add_normal(
    '기쁨은 환경에서 오지 않습니다. 빌립보서를 쓴 바울은 감옥에 있었습니다. '
    '자유도 없고, 편안함도 없고, 내일도 불확실했습니다. '
    '그런데 그는 "주 안에서 항상 기뻐하라"라고 썼습니다. '
    '이것은 환경을 초월하는 기쁨이 가능하다는 증거입니다.'
)

add_sub('감사와 기쁨의 실천')
add_bullet('감사 일기: 매일 감사한 것 3가지를 적어봅니다. 한 달만 지속하면 감사의 눈이 열리기 시작합니다.')
add_bullet('찬미의 시간: 어려울 때일수록 찬양하십시오. 찬양은 마음의 분위기를 바꾸는 강력한 도구입니다.')
add_bullet('불평 대신 감사로 바꾸기: "일이 너무 많다" → "일할 수 있는 능력을 주셔서 감사합니다." 이것은 현실을 부정하는 것이 아닙니다. 같은 현실을 하나님의 관점에서 다시 보는 것입니다.')

# ── 원리 4 ──
add_ref('관련 설교: "보혈로 그린 자화상", "건전한 자존심", "나는 누구인가", "삶을 좀더 밝게 보며", "마음의 눈", "보라 새 것이 되었도다"')

add_normal('원리 4. 새 자아상을 입어라 — 하나님의 눈으로 나를 보기', bold=True, size=13, space_before=16, space_after=8)

add_sub('세상의 자화상 vs 보혈의 자화상')
add_normal(
    '생각을 바꾸는 가장 근본적인 작업은 \'나는 누구인가\'에 대한 답을 바꾸는 것입니다.'
)
blank()

add_normal('세상은 끊임없이 우리에게 자화상을 그려줍니다:')
add_bullet('성적, 학벌, 외모, 재산으로 사람의 가치를 매깁니다.')
add_bullet('실패하면 \'실패자\'라는 딱지를 붙입니다.')
add_bullet('비교하고, 경쟁하고, 부족함을 느끼게 합니다.')

blank()
add_normal('그러나 하나님은 전혀 다른 자화상을 보여주십니다:')
add_bullet('"너는 내 눈에 보배롭고 존귀하며 내가 너를 사랑하노라." (사 43:4)')
add_bullet('"너희는 택하신 족속이요 왕 같은 제사장들이요 거룩한 나라요 그의 소유가 된 백성이니." (벧전 2:9)')
add_bullet('"누구든지 그리스도 안에 있으면 새로운 피조물이라." (고후 5:17)')

add_sub('자아상이 바뀌면 생각이 바뀐다')
add_normal(
    '나를 \'실패자\'로 보면, 모든 상황이 실패의 증거로 보입니다. '
    '나를 \'하나님의 자녀\'로 보면, 어려운 상황도 성장의 과정으로 보입니다. '
    '자아상은 생각의 필터입니다. 이 필터를 바꾸면 같은 세상이 완전히 다르게 보이기 시작합니다.'
)

add_sub('새 자아상을 세우는 방법')
add_bullet('하나님이 나에 대해 말씀하신 것을 찾기: 성경에서 \'하나님의 자녀\'가 누구인지 찾아 읽으십시오. 그것이 바로 당신의 참 모습입니다.')
add_bullet('매일 고백하기: "나는 하나님의 자녀입니다. 나는 그리스도 안에서 새로운 피조물입니다. 나는 보배롭고 존귀한 존재입니다." 이것은 하나님의 말씀에 근거한 진실의 고백입니다.')
add_bullet('과거의 실패에서 벗어나기: 이전 것은 지나갔습니다. 실패의 기억이 떠오를 때, "보라 새 것이 되었도다"를 묵상하십시오.')

# ══════════════════════════════════════════════════════════
# 나가며
# ══════════════════════════════════════════════════════════
doc.add_heading('나가며', level=1)
doc.add_heading('생각이 바뀌면 삶이 바뀐다', level=2)

add_normal(
    '이 네 가지 원리는 한 번의 결심이 아니라 매일의 훈련입니다. '
    '농부가 매일 밭을 가꾸듯이, 우리도 매일 마음의 밭을 가꿔야 합니다.'
)
blank()

add_quote('"우리가 낙심하지 아니하노니 우리의 겉사람은 낡아지나 우리의 속사람은 날로 새로워지도다." (고린도후서 4:16)')
blank()

add_normal(
    '생각을 바꾸는 것은 하루아침에 되지 않습니다. '
    '그러나 오늘 한 걸음을 시작하면, 내일은 어제보다 한 걸음 앞서 있을 것입니다. '
    '하나님은 우리의 한 걸음 한 걸음을 보시며 함께 걸어주십니다.'
)
blank()

add_quote('"무릇 진실하며 무릇 존귀하며 무릇 옳으며 무릇 정결하며 무릇 사랑 받을 만하며 무릇 칭찬 받을 만하며 무슨 덕이 있든지 무슨 기림이 있든지 이것들을 생각하라." (빌립보서 4:8)')

# ── 정리표 ──
blank()
add_normal('[ 정리 ] 생각을 바꾸는 4가지 실천 원리', bold=True, size=12, space_before=12)
blank()

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['원리', '핵심', '방향']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

data = [
    ['원리 1. 비워라', '염려, 원망, 미움, 열등감을 내려놓기', '마음의 청소'],
    ['원리 2. 심어라', '하나님의 말씀을 읽고, 암송하고, 묵상하기', '생각의 프로그래밍'],
    ['원리 3. 채워라', '감사와 기쁨으로 마음의 체질을 바꾸기', '마음의 건강'],
    ['원리 4. 입어라', '보혈로 그린 새 자아상을 받아들이기', '정체성의 전환'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

# ── 참고 설교 ──
blank()
add_normal('[ 참고 ] 본 아티클에서 인용된 주요 설교', bold=True, size=10, color=RGBColor(0x66,0x66,0x66), space_before=16)
refs = [
    '"행복과 마음의 작정" (1981.1.11)',
    '"삶을 실패케 하는 마음가짐" (1982.1.24)',
    '"마음의 청소" (1982.7.4)',
    '"생각을 바꿔라" (1984.11.11)',
    '"마음의 건강" (1985.9.22)',
    '"네 마음을 지키라" (1986.1.19)',
    '"마음의 변화" (1988.6.12)',
    '"마음의 눈" (1989.12.3)',
    '"마음의 평강" (1989.7.30)',
    '"삶을 좀더 밝게 보며" (1990.1.28)',
    '"보혈로 그린 자화상" (1992.3.1)',
    '"축복을 가져오는 생활태도" (1996.10.6)',
]
for ref in refs:
    add_ref(f'  - {ref}')

# ── 저장 ──
output_path = 'output/4차원영성_생각_아티클.docx'
doc.save(output_path)
print(f'✓ 생성 완료: {output_path}')
